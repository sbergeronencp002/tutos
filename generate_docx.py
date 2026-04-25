import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

TUTO_FILES = [f"/home/user/tutos/tuto_{i}.md" for i in range(1, 16)]

EMOJI_MAP = {
    "🎯": "🎯", "🚀": "🚀", "✨": "✨", "💡": "💡", "🖥️": "🖥️",
    "💾": "💾", "👀": "👀", "🎉": "🎉", "🧠": "🧠", "❓": "❓",
    "➡️": "➡️", "🧹": "🧹", "✏️": "✏️", "🔬": "🔬", "✅": "✅",
    "🟡": "🟡", "🔵": "🔵", "🤔": "🤔", "🔔": "🔔", "🔄": "🔄",
    "❤️": "❤️", "🤍": "🤍", "🔁": "🔁", "➕": "➕", "➖": "➖",
    "⏱️": "⏱️", "⚔️": "⚔️", "🛡️": "🛡️", "🏹": "🏹", "🎨": "🎨",
    "🪨": "🪨", "📄": "📄", "✂️": "✂️", "📡": "📡", "⚠️": "⚠️",
    "🌞": "🌞", "🌙": "🌙", "❄️": "❄️", "📐": "📐", "🔒": "🔒",
    "🎱": "🎱", "🏆": "🏆", "⬆️": "⬆️", "⬇️": "⬇️", "⬅️": "⬅️",
    "ℹ️": "ℹ️",
}

def strip_inline_md(text):
    """Remove bold/italic markdown and backtick formatting for plain text."""
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    text = re.sub(r'``\|[^|]+:[^|]+\|(.+?)\|\|``', r'\1', text)
    text = re.sub(r'`([^`]+)`', r'\1', text)
    return text

def add_heading(doc, text, level):
    text = strip_inline_md(text)
    p = doc.add_heading(text, level=level)
    p.style.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D) if level == 1 else RGBColor(0x2E, 0x74, 0xB5)
    return p

def add_code_block(doc, code):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(code.strip())
    run.font.name = "Courier New"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x23, 0x23, 0x23)
    # Light grey background via shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F0F0F0')
    pPr.append(shd)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    return p

def add_blockquote(doc, text):
    text = strip_inline_md(text)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    run.font.italic = True
    run.font.size = Pt(10)
    return p

def add_body(doc, text):
    text = strip_inline_md(text)
    p = doc.add_paragraph(text)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(3)
    for run in p.runs:
        run.font.size = Pt(11)
    return p

def add_page_break(doc):
    doc.add_page_break()

def parse_and_add(doc, md_text):
    lines = md_text.splitlines()
    i = 0
    in_code = False
    code_lines = []

    while i < len(lines):
        line = lines[i]

        # Code block start/end
        if line.strip().startswith("```"):
            if not in_code:
                in_code = True
                code_lines = []
            else:
                in_code = False
                add_code_block(doc, "\n".join(code_lines))
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        # Headings
        if line.startswith("# ") and not line.startswith("## "):
            add_heading(doc, line[2:], 1)
        elif line.startswith("## @showdialog"):
            pass  # skip directive
        elif line.startswith("## "):
            add_heading(doc, line[3:], 2)
        elif line.startswith("> "):
            add_blockquote(doc, line[2:])
        elif line.strip() == "":
            pass  # skip blank lines
        else:
            add_body(doc, line)

        i += 1

def main():
    doc = Document()

    # Document title
    title = doc.add_heading("Tutoriels micro:bit — 1 à 15", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Style defaults
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    for idx, filepath in enumerate(TUTO_FILES):
        if idx > 0:
            add_page_break(doc)
        with open(filepath, encoding="utf-8") as f:
            content = f.read()
        parse_and_add(doc, content)

    out = "/home/user/tutos/tutoriels_microbit.docx"
    doc.save(out)
    print(f"Saved: {out}")

if __name__ == "__main__":
    main()
